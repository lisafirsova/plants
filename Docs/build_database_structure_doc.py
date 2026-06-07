from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Описание_структуры_БД_Plants.docx")


TABLES = [
    (
        "app_users",
        "пользователях мобильного приложения",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор пользователя, первичный ключ таблицы, формируется автоматически."),
            ("google_id", "VARCHAR(200)", "до 200 символов", "Уникальный идентификатор Google-аккаунта, используемый для авторизации пользователя."),
            ("name", "VARCHAR(160)", "до 160 символов", "Имя пользователя, полученное из профиля Google."),
            ("email", "VARCHAR(200)", "до 200 символов", "Адрес электронной почты пользователя."),
            ("avatar_url", "TEXT", "переменный", "Ссылка на изображение профиля пользователя."),
            ("is_admin", "BOOLEAN", "1 Б", "Признак наличия прав администратора. Ограничение базы данных допускает только одного администратора."),
        ],
    ),
    (
        "app_plant_types",
        "категориях растений",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор категории растения, первичный ключ таблицы."),
            ("name", "VARCHAR(40)", "до 40 символов", "Уникальное название категории: «Лиственные», «Цветущие» или «Суккуленты»."),
        ],
    ),
    (
        "app_plant_species",
        "видах растений",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор вида растения, первичный ключ таблицы."),
            ("plant_type_id", "INTEGER", "4 Б", "Идентификатор категории растения, внешний ключ на таблицу «app_plant_types»."),
            ("name", "VARCHAR(120)", "до 120 символов", "Русскоязычное название вида растения."),
            ("latin_name", "VARCHAR(160)", "до 160 символов", "Латинское наименование вида растения."),
            ("care_description", "TEXT", "переменный", "Общие рекомендации по содержанию и уходу за растением данного вида."),
        ],
    ),
    (
        "app_period_units",
        "единицах измерения периодичности ухода",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор единицы периода, первичный ключ таблицы."),
            ("code", "VARCHAR(20)", "до 20 символов", "Уникальный программный код единицы периода: day, week или month."),
            ("name", "VARCHAR(40)", "до 40 символов", "Отображаемое название единицы периода: день, неделя или месяц."),
            ("day_multiplier", "INTEGER", "4 Б", "Коэффициент перевода выбранной единицы периода в количество дней."),
        ],
    ),
    (
        "app_plants",
        "карточках растений пользователей",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор растения, первичный ключ таблицы."),
            ("user_id", "INTEGER", "4 Б", "Идентификатор владельца растения, внешний ключ на таблицу «app_users»."),
            ("plant_type_id", "INTEGER", "4 Б", "Идентификатор категории растения, внешний ключ на таблицу «app_plant_types»."),
            ("species_id", "INTEGER", "4 Б", "Идентификатор вида растения, внешний ключ на таблицу «app_plant_species»; может быть не указан."),
            ("name", "VARCHAR(120)", "до 120 символов", "Пользовательское название растения."),
            ("type", "VARCHAR(40)", "до 40 символов", "Текстовое представление категории растения, сохранённое для совместимости с мобильным клиентом."),
            ("watering_enabled", "BOOLEAN", "1 Б", "Признак наличия активного расписания полива."),
            ("fertilizer_enabled", "BOOLEAN", "1 Б", "Признак наличия активного расписания подкормки."),
            ("lifecycle_phase", "VARCHAR(80)", "до 80 символов", "Текущая фаза жизненного цикла растения."),
            ("health_status", "VARCHAR(80)", "до 80 символов", "Текущее состояние растения."),
            ("location", "VARCHAR(160)", "до 160 символов", "Место размещения растения."),
            ("notes", "TEXT", "переменный", "Пользовательские заметки о растении и особенностях ухода."),
        ],
    ),
    (
        "app_watering_schedules",
        "расписаниях мероприятий по уходу",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор расписания, первичный ключ таблицы."),
            ("plant_id", "INTEGER", "4 Б", "Идентификатор растения, внешний ключ на таблицу «app_plants»."),
            ("period_unit_id", "INTEGER", "4 Б", "Идентификатор единицы периодичности, внешний ключ на таблицу «app_period_units»."),
            ("start_date", "DATE", "4 Б", "Дата начала отсчёта расписания."),
            ("period_days", "INTEGER", "4 Б", "Рассчитанная продолжительность периода между мероприятиями в днях."),
            ("period_value", "INTEGER", "4 Б", "Числовое значение выбранной пользователем периодичности."),
            ("notification_time", "TIME", "8 Б", "Время отправки локального уведомления."),
            ("type", "VARCHAR(20)", "до 20 символов", "Тип мероприятия: Watering, Fertilizer, Pruning или Repotting."),
            ("last_completed_date", "DATE", "4 Б", "Дата последнего выполнения мероприятия; может отсутствовать."),
        ],
    ),
    (
        "app_care_events",
        "конкретных событиях и результатах выполнения мероприятий по уходу",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор события ухода, первичный ключ таблицы."),
            ("schedule_id", "INTEGER", "4 Б", "Идентификатор расписания, внешний ключ на таблицу «app_watering_schedules»."),
            ("plant_id", "INTEGER", "4 Б", "Идентификатор растения, внешний ключ на таблицу «app_plants»."),
            ("scheduled_date", "DATE", "4 Б", "Плановая дата выполнения мероприятия."),
            ("completed_at", "TIMESTAMP", "8 Б", "Фактические дата и время выполнения; могут отсутствовать."),
            ("status", "VARCHAR(20)", "до 20 символов", "Статус задачи: Planned, Done или Skipped."),
            ("note", "TEXT", "переменный", "Заметка пользователя о выполненном или пропущенном мероприятии."),
        ],
    ),
    (
        "app_photos",
        "фотографиях растений в фотоархиве",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор фотографии, первичный ключ таблицы."),
            ("plant_id", "INTEGER", "4 Б", "Идентификатор растения, внешний ключ на таблицу «app_plants»."),
            ("title", "VARCHAR(160)", "до 160 символов", "Пользовательское название фотографии."),
            ("file_path", "TEXT", "переменный", "Путь к файлу изображения в хранилище приложения."),
            ("date_taken", "TIMESTAMP", "8 Б", "Дата и время создания либо добавления фотографии."),
        ],
    ),
    (
        "app_reference_categories",
        "категориях записей справочника",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор категории справочника, первичный ключ таблицы."),
            ("code", "VARCHAR(20)", "до 20 символов", "Уникальный программный код категории: pest или disease."),
            ("name", "VARCHAR(80)", "до 80 символов", "Название категории: «Вредители» или «Болезни»."),
        ],
    ),
    (
        "app_reference_items",
        "вредителях и болезнях, представленных в справочнике",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор записи справочника, первичный ключ таблицы."),
            ("category_id", "INTEGER", "4 Б", "Идентификатор категории, внешний ключ на таблицу «app_reference_categories»."),
            ("name", "VARCHAR(160)", "до 160 символов", "Название вредителя или болезни."),
            ("treatment_description", "TEXT", "переменный", "Описание признаков, профилактики и рекомендуемого способа лечения."),
            ("image_path", "TEXT", "переменный", "Путь к изображению вредителя или проявления болезни."),
        ],
    ),
    (
        "app_reference_item_plant_types",
        "соответствии записей справочника категориям растений",
        [
            ("reference_item_id", "INTEGER", "4 Б", "Идентификатор записи справочника, внешний ключ на таблицу «app_reference_items» и часть составного первичного ключа."),
            ("plant_type_id", "INTEGER", "4 Б", "Идентификатор категории растения, внешний ключ на таблицу «app_plant_types» и часть составного первичного ключа."),
        ],
    ),
    (
        "app_issue_observations",
        "выявленных у растений вредителях и заболеваниях",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор наблюдения, первичный ключ таблицы."),
            ("plant_id", "INTEGER", "4 Б", "Идентификатор растения, внешний ключ на таблицу «app_plants»."),
            ("reference_item_id", "INTEGER", "4 Б", "Идентификатор выявленной болезни или вредителя, внешний ключ на таблицу «app_reference_items»."),
            ("observed_at", "DATE", "4 Б", "Дата обнаружения проблемы."),
            ("note", "TEXT", "переменный", "Дополнительное описание состояния растения."),
        ],
    ),
    (
        "app_recommendation_settings",
        "параметрах автоматического формирования рекомендаций по уходу",
        [
            ("code", "VARCHAR(60)", "до 60 символов", "Уникальный код настройки, первичный ключ таблицы."),
            ("numeric_value", "NUMERIC(10,2)", "переменный", "Числовое значение параметра алгоритма рекомендаций."),
            ("description", "VARCHAR(240)", "до 240 символов", "Назначение и текстовое описание настройки."),
        ],
    ),
    (
        "app_protection_products",
        "средствах защиты и лечения растений",
        [
            ("id", "INTEGER", "4 Б", "Уникальный идентификатор средства защиты, первичный ключ таблицы."),
            ("name", "VARCHAR(160)", "до 160 символов", "Уникальное название средства защиты."),
            ("active_ingredient", "VARCHAR(200)", "до 200 символов", "Действующее вещество средства."),
            ("application_description", "TEXT", "переменный", "Описание порядка и особенностей применения."),
            ("hazard_class", "VARCHAR(80)", "до 80 символов", "Класс или уровень опасности средства."),
        ],
    ),
]


def set_font(run, name="Times New Roman", size=14, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_fixed_layout(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width_twips = int(widths_cm[index] / 2.54 * 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_twips))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def format_body_paragraph(paragraph, first_line=True):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    for run in paragraph.runs:
        set_font(run)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run)
    format_body_paragraph(paragraph)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph


def add_structure_table(doc, number, name, purpose, rows):
    introduction = add_body(
        doc,
        f"Таблица «{name}» хранит информацию о {purpose}, структура которой представлена в таблице {number}.",
    )
    introduction.paragraph_format.keep_with_next = True
    add_caption(doc, f"Таблица {number} – Структура таблицы «{name}»")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ("Имя поля", "Тип данных", "Размер поля", "Описание")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "D9E2F3")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                set_font(run, size=11, bold=True)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])

    for field_name, data_type, field_size, description in rows:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        values = (field_name, data_type, field_size, description)
        for index, value in enumerate(values):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index in (0, 3) else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_font(run, size=10)

    set_table_fixed_layout(table, [3.1, 3.2, 3.0, 7.2])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_font(run, size=12)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_page_number(section.footer.paragraphs[0])

    opening = [
        (
            "Реляционная модель хранения данных позволяет чётко организовать информацию и установить связи между объектами предметной области. "
            "В реляционной модели данные хранятся в отношениях, представленных в виде двумерных таблиц. Строки таблицы соответствуют отдельным записям, "
            "а столбцы – атрибутам сущностей. Каждая запись должна однозначно идентифицироваться. Для этого применяются первичные ключи, а для установления "
            "связей между таблицами используются внешние ключи."
        ),
        (
            "Физическая модель базы данных мобильного приложения «Plants» реализована в системе управления базами данных PostgreSQL. "
            "Структура базы данных включает четырнадцать таблиц, предназначенных для хранения сведений о пользователях, растениях, расписаниях ухода, "
            "фотоархиве, вредителях, болезнях, средствах защиты и настройках автоматических рекомендаций."
        ),
        (
            "Структура таблиц базы данных, назначение их полей и используемые типы данных приведены в таблицах 3.1–3.14."
        ),
    ]
    for text in opening:
        add_body(doc, text)

    for index, (name, purpose, rows) in enumerate(TABLES, start=1):
        add_structure_table(doc, f"3.{index}", name, purpose, rows)

    add_body(
        doc,
        "Физическая схема базы данных представлена на рисунке 3.1.",
    )
    placeholder = doc.add_paragraph()
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder.paragraph_format.space_before = Pt(18)
    placeholder.paragraph_format.space_after = Pt(18)
    run = placeholder.add_run("[Вставить физическую схему базы данных]")
    set_font(run, size=12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    figure_caption = doc.add_paragraph()
    figure_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_caption.paragraph_format.space_after = Pt(6)
    run = figure_caption.add_run("Рисунок 3.1 – Физическая схема базы данных мобильного приложения «Plants»")
    set_font(run, size=14)

    closing = [
        (
            "Авторизация пользователей выполняется с использованием Google OAuth. Пароли пользователей и токены доступа в базе данных приложения не хранятся. "
            "Для идентификации учётной записи сохраняется уникальный идентификатор Google-аккаунта, а разграничение прав выполняется посредством логического поля "
            "«is_admin». Частичный уникальный индекс «app_users_single_admin_idx» гарантирует, что права администратора могут быть назначены только одной учётной записи."
        ),
        (
            "Целостность данных обеспечивается первичными ключами, ограничениями уникальности, проверочными ограничениями и внешними ключами. "
            "При удалении пользователя каскадно удаляются принадлежащие ему растения. При удалении растения удаляются связанные расписания, события ухода, фотографии "
            "и зафиксированные наблюдения. Удаление записи справочника приводит к удалению её связей с категориями растений и наблюдениями. Для справочных сущностей, "
            "удаление которых может нарушить существующие карточки растений, применяются ограничения RESTRICT либо SET NULL."
        ),
        (
            "Таблицы категорий растений, видов растений, единиц периодичности и категорий справочника вынесены в отдельные отношения, что уменьшает дублирование данных. "
            "Связь многие-ко-многим между записями справочника и категориями растений реализована через промежуточную таблицу «app_reference_item_plant_types». "
            "Такое построение соответствует принципам нормализации и позволяет независимо изменять справочные данные без нарушения связанных записей."
        ),
        (
            "Для защиты от SQL-инъекций взаимодействие приложения с PostgreSQL выполняется с помощью параметризованных запросов Npgsql. "
            "Ключ доступа к сервису искусственного интеллекта не сохраняется в мобильном приложении и используется только серверным компонентом AI Proxy."
        ),
    ]
    for text in closing:
        add_body(doc, text)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
