from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT = Path(__file__).with_name("Тестирование_приложения_Plants_исправлено.docx")


TEST_CASES = [
    {
        "number": "3.15",
        "precondition": (
            "Предварительное условие для проверки мобильного приложения со стороны пользователя: "
            "необходимо запустить приложение Plants на Android-устройстве или эмуляторе. "
            "Окно авторизации представлено в приложении В на рисунке В.1."
        ),
        "intro": "Тест-кейс, описывающий процесс авторизации пользователя, приведён в таблице 3.15.",
        "module": "Авторизация пользователя",
        "steps": [
            "Запустить мобильное приложение Plants.",
            "Нажать кнопку «Войти через Google».",
            "Выбрать Google-аккаунт в системном окне.",
            "Подтвердить предоставление приложению доступа к основным данным профиля.",
            "Дождаться завершения входа и открытия главного экрана приложения.",
        ],
        "expected": (
            "пользователь успешно авторизуется, данные Google-аккаунта сохраняются в базе данных, "
            "окно входа закрывается и открывается экран «Полив». При повторном запуске приложения "
            "авторизация не запрашивается до тех пор, пока пользователь самостоятельно не выйдет из аккаунта."
        ),
        "figures": [
            "Окно авторизации представлено на рисунке В.1.",
            "Системное окно выбора Google-аккаунта представлено на рисунке В.2.",
        ],
        "actual": (
            "Фактический результат соответствует ожидаемому. Пользователь вошёл в систему, сведения об учётной "
            "записи были сохранены в таблице «app_users», после чего открылся экран «Полив». "
            "Результат представлен на рисунке В.3."
        ),
    },
    {
        "number": "3.16",
        "precondition": (
            "Предварительное условие для проверки добавления растения: пользователь должен быть авторизован "
            "в системе. Тест-кейс, описывающий процесс добавления растения, приведён в таблице 3.16."
        ),
        "intro": None,
        "module": "Добавление растения",
        "steps": [
            "Открыть вкладку «Растения».",
            "Выбрать категорию «Лиственные».",
            "Нажать кнопку «Добавить».",
            "Ввести значение «Монстера» в поле «Имя».",
            "Выбрать вид растения и фазу жизненного цикла.",
            "Включить параметры «Полив» и «Удобрения».",
            "Указать даты начала, периодичность и время уведомлений.",
            "Нажать кнопку «Добавить растение».",
            "Перейти на вкладку «Полив» и проверить появление мероприятий по уходу.",
        ],
        "expected": (
            "карточка растения сохраняется в базе данных и связывается с текущим пользователем. "
            "Для включённых мероприятий создаются расписания полива и подкормки, а растение появляется "
            "в календаре ухода и фотоотчёте."
        ),
        "figures": [
            "Страница выбора категории растения представлена на рисунке В.4.",
            "Заполненная форма добавления растения представлена на рисунке В.5.",
        ],
        "actual": (
            "Фактический результат соответствует ожидаемому. Растение было добавлено в таблицу «app_plants», "
            "расписания сохранены в таблице «app_watering_schedules», а соответствующие задачи появились "
            "на вкладке «Полив». Результат представлен на рисунке В.6."
        ),
    },
    {
        "number": "3.17",
        "precondition": (
            "Предварительное условие для проверки изменения статуса задачи ухода: пользователь должен быть "
            "авторизован, а для одного из его растений должно быть создано расписание полива или подкормки. "
            "Тест-кейс, описывающий изменение статуса задачи ухода, приведён в таблице 3.17."
        ),
        "intro": None,
        "module": "Изменение статуса задачи ухода",
        "steps": [
            "Открыть вкладку «Полив».",
            "Найти задачу, запланированную на текущую дату.",
            "Нажать кнопку «Полить» или «Удобрить».",
            "Проверить изменение текста и оформления кнопки.",
            "Повторно нажать кнопку выполненного действия.",
            "Проверить возврат задачи в невыполненное состояние.",
        ],
        "expected": (
            "после первого нажатия задача получает статус выполненной, кнопка изменяется на «Полито» или "
            "«Удобрено», а результат сохраняется в базе данных. После повторного нажатия выполнение отменяется, "
            "запись о событии удаляется и кнопка возвращается в исходное состояние."
        ),
        "figures": [
            "Задача ухода до выполнения представлена на рисунке В.7.",
        ],
        "actual": (
            "Фактический результат соответствует ожидаемому. Статус задачи был сохранён в таблице "
            "«app_care_events», кнопка изменила состояние без перезапуска экрана. После повторного нажатия "
            "выполнение было отменено. Результат представлен на рисунке В.8."
        ),
    },
    {
        "number": "3.18",
        "precondition": (
            "Предварительное условие для проверки фотоархива: пользователь должен быть авторизован и иметь "
            "хотя бы одно растение. Тест-кейс, описывающий процесс добавления фотографии растения, "
            "приведён в таблице 3.18."
        ),
        "intro": None,
        "module": "Добавление фотографии в фотоархив",
        "steps": [
            "Открыть вкладку «Полив».",
            "Переключить сегмент на значение «Фотоотчёт».",
            "Нажать кнопку «Сделать фото» напротив выбранного растения.",
            "Нажать кнопку «Добавить фото».",
            "Выбрать или сделать фотографию растения.",
            "Ввести название фотографии.",
            "Указать дату фотографии.",
            "Нажать кнопку «Создать».",
            "Проверить появление фотографии в фотоархиве растения.",
        ],
        "expected": (
            "выбранное изображение сохраняется в хранилище приложения, а его название, путь к файлу, дата "
            "и идентификатор растения записываются в базу данных. Фотография отображается в фотоархиве "
            "выбранного растения."
        ),
        "figures": [
            "Список растений в режиме «Фотоотчёт» представлен на рисунке В.9.",
            "Форма добавления фотографии представлена на рисунке В.10.",
        ],
        "actual": (
            "Фактический результат соответствует ожидаемому. Файл изображения был сохранён, запись добавлена "
            "в таблицу «app_photos», а фотография с указанным названием и датой появилась в фотоархиве. "
            "Результат представлен на рисунке В.11."
        ),
    },
    {
        "number": "3.19",
        "precondition": (
            "Предварительное условие для проверки поиска и фильтрации справочника: необходимо запустить "
            "приложение и открыть вкладку «Справочник». Страница справочника представлена в приложении В "
            "на рисунке В.12. Тест-кейс, описывающий процесс поиска и фильтрации вредителей и болезней, "
            "приведён в таблице 3.19."
        ),
        "intro": None,
        "module": "Поиск и фильтрация справочника",
        "steps": [
            "Открыть вкладку «Справочник».",
            "Нажать на иконку поиска.",
            "Ввести значение «гниль» в поле поиска.",
            "Нажать на иконку фильтрации.",
            "Выбрать категорию «Болезни».",
            "Выбрать тип растения «Суккуленты».",
            "Нажать кнопку «Применить».",
            "Открыть список болезней и проверить отображаемые записи.",
        ],
        "expected": (
            "справочник отображает только записи, соответствующие поисковому запросу, выбранной категории "
            "и типу растения. При открытии подходящей записи отображаются её изображение и рекомендации по лечению."
        ),
        "figures": [
            "Страница справочника до применения поиска представлена на рисунке В.12.",
            "Панель фильтрации с выбранными параметрами представлена на рисунке В.13.",
        ],
        "actual": (
            "Фактический результат соответствует ожидаемому. После применения параметров в списке остались "
            "только соответствующие болезни для суккулентов, содержащие введённый текст. "
            "Результат представлен на рисунке В.14."
        ),
    },
    {
        "number": "3.20",
        "precondition": (
            "Предварительное условие для проверки работы ИИ-консультанта: пользователь должен быть авторизован, "
            "мобильное устройство должно иметь доступ к сети, а сервер Plants.AiProxy должен быть запущен "
            "с настроенным ключом OpenAI API. Тест-кейс, описывающий получение рекомендации по фотографии, "
            "приведён в таблице 3.20."
        ),
        "intro": None,
        "module": "Получение рекомендации ИИ-консультанта",
        "steps": [
            "Открыть вкладку «Справочник».",
            "Нажать карточку «ИИ-консультант».",
            "Нажать кнопку «Прикрепить фото».",
            "Выбрать фотографию растения.",
            "Ввести вопрос о состоянии растения.",
            "Нажать кнопку «Отправить».",
            "Дождаться ответа ИИ-консультанта.",
            "Проверить отображение рекомендаций в истории диалога.",
        ],
        "expected": (
            "приложение передаёт вопрос, фотографию, сведения о растениях пользователя и историю диалога "
            "серверному компоненту. ИИ анализирует видимые признаки и возвращает рекомендации о возможной "
            "болезни, вредителе, поливе, обрезке или необходимости пересадки."
        ),
        "figures": [
            "Окно ИИ-консультанта представлено на рисунке В.15.",
            "Сообщение с прикреплённой фотографией представлено на рисунке В.16.",
        ],
        "actual": (
            "Фактический результат соответствует ожидаемому. Вопрос и изображение были отправлены серверу, "
            "после чего в чате отобразился ответ с описанием видимых признаков, возможных причин и безопасных "
            "рекомендаций по дальнейшему уходу. Результат представлен на рисунке В.17."
        ),
    },
    {
        "number": "3.21",
        "precondition": (
            "Предварительное условие для проверки административной функции: необходимо авторизоваться под "
            "единственной учётной записью с правами администратора. Тест-кейс, описывающий добавление записи "
            "в справочник администратором, приведён в таблице 3.21."
        ),
        "intro": None,
        "module": "Добавление записи справочника администратором",
        "steps": [
            "Открыть панель администратора.",
            "Нажать кнопку «Добавить».",
            "Ввести название вредителя или болезни.",
            "Ввести описание способа лечения.",
            "Выбрать изображение.",
            "Выбрать категорию «Вредитель» или «Болезнь».",
            "Отметить подходящие типы растений: лиственные, цветущие или суккуленты.",
            "Нажать кнопку «Создать».",
            "Открыть справочник и проверить наличие новой записи.",
        ],
        "expected": (
            "новая запись сохраняется в базе данных с указанной категорией, описанием, изображением и типами "
            "растений. После сохранения запись становится доступна в справочнике и учитывается при фильтрации."
        ),
        "figures": [
            "Панель администратора представлена на рисунке В.18.",
            "Заполненная форма добавления записи представлена на рисунке В.19.",
        ],
        "actual": (
            "Фактический результат соответствует ожидаемому. Запись была добавлена в таблицу "
            "«app_reference_items», а связи с типами растений сохранены в таблице "
            "«app_reference_item_plant_types». Новая запись появилась в справочнике. "
            "Результат представлен на рисунке В.20."
        ),
    },
    {
        "number": "3.22",
        "precondition": (
            "Предварительное условие для проверки формирования отчётности: необходимо авторизоваться под "
            "учётной записью администратора и открыть раздел отчётов. В базе данных должны находиться сведения "
            "о растениях и мероприятиях по уходу. Тест-кейс, описывающий формирование и экспорт отчёта, "
            "приведён в таблице 3.22."
        ),
        "intro": None,
        "module": "Формирование и экспорт отчёта",
        "steps": [
            "Открыть панель администратора.",
            "Нажать кнопку «Экспорт».",
            "Выбрать отчётный месяц.",
            "Сформировать план работ на месяц.",
            "Проверить отображение мероприятий по уходу.",
            "Сформировать статистический отчёт по растениям, вредителям или болезням.",
            "Нажать кнопку экспорта.",
            "Выбрать приложение для сохранения или отправки сформированного файла.",
        ],
        "expected": (
            "система получает данные из PostgreSQL, формирует выбранные аналитические таблицы и создаёт файл "
            "отчёта, доступный для сохранения или передачи через системное меню Android."
        ),
        "figures": [
            "Раздел формирования отчётов представлен на рисунке В.21.",
            "Сформированный отчёт за выбранный период представлен на рисунке В.22.",
        ],
        "actual": (
            "Фактический результат соответствует ожидаемому. Отчёт был сформирован на основании актуальных "
            "данных пользователя, после чего открылось системное меню Android для сохранения или отправки файла. "
            "Результат представлен на рисунке В.23."
        ),
    },
]


def set_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def format_paragraph(paragraph, *, justify=True, indent=True, keep_next=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = keep_next
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)


def add_text(doc, text, *, keep_next=False):
    paragraph = doc.add_paragraph()
    format_paragraph(paragraph, keep_next=keep_next)
    set_font(paragraph.add_run(text))
    return paragraph


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def prevent_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(element)


def set_table_widths(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = int(widths_cm[index] / 2.54 * 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def fill_cell(cell, text, *, size=11, bold=False, center=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    set_font(paragraph.add_run(text), size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_cell_paragraph(cell, text, *, size=11, bold=False, first=False):
    paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
    if first:
        paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    set_font(paragraph.add_run(text), size=size, bold=bold)
    return paragraph


def add_test_case(doc, case):
    add_text(doc, case["precondition"], keep_next=True)
    if case["intro"]:
        add_text(doc, case["intro"], keep_next=True)

    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(3)
    caption.paragraph_format.keep_with_next = True
    set_font(caption.add_run(f"Таблица {case['number']} – Тест-кейс"))

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Модуль / функция", "Шаги воспроизведения", "Результат"]
    for index, header in enumerate(headers):
        fill_cell(table.rows[0].cells[index], header, size=11, bold=True, center=True)
        shade(table.rows[0].cells[index], "D9E2F3")
    repeat_header(table.rows[0])
    prevent_split(table.rows[0])

    data_row = table.add_row()
    prevent_split(data_row)
    fill_cell(data_row.cells[0], case["module"], size=11, center=True)
    fill_cell(
        data_row.cells[1],
        "\n".join(f"{index}. {step}" for index, step in enumerate(case["steps"], start=1)),
        size=11,
    )

    result_cell = data_row.cells[2]
    result_cell.text = ""
    result_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_cell_paragraph(
        result_cell,
        f"Ожидаемый: {case['expected']}",
        size=11,
        first=True,
    )
    for figure in case["figures"]:
        add_cell_paragraph(result_cell, figure, size=11)
    add_cell_paragraph(result_cell, case["actual"], size=11)

    set_table_widths(table, [4.0, 7.1, 5.2])

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_font(run, size=12)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_page_number(section.footer.paragraphs[0])

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for case in TEST_CASES:
        add_test_case(doc, case)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
